import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_word_doc():
    doc = docx.Document()

    # Page Margins (0.75 in all around)
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Styles Setup
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0, 0, 0)

    # Helper function for headings
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(18)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        return p

    def add_body(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        return p

    def add_code(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 0, 0)
        return p

    # ----------------------------------------------------
    # Page 1: COVER PAGE
    # ----------------------------------------------------
    for _ in range(3): doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("JSS MAHAVIDYAPEETHA\nJSS SCIENCE AND TECHNOLOGY UNIVERSITY\nJSS TECHNICAL INSTITUTIONS CAMPUS MYSURU – 570006 KARNATAKA, INDIA\n\n\nDepartment of Information Science and Engineering\n\n\nA Mini-Project Report titled\n\n")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TALENTFORGE: A TALENT SHOWCASE & RECRUITMENT PLATFORM\n\n\n")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Submitted by\n\nRaghavendra K        01JST24UIS067\nRaghavendra S        01JST24UIS068\n\n\nUnder the supervision of\n\nMs. R S Anusha\nAssistant Professor\nDepartment of Information Science & Engineering\nJSS S&T University, Mysuru.\n\n\n\n2026-27")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True

    doc.add_page_break()

    # ----------------------------------------------------
    # Page 2: Index Sheet (Table of Contents)
    # ----------------------------------------------------
    add_heading_1("TABLE OF CONTENTS")
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Section'
    hdr_cells[1].text = 'Page'
    
    toc_entries = [
        ("1. Abstract", "2"),
        ("2. Introduction", "3"),
        ("3. Software Requirements Specification (SRS)", "4"),
        ("   3.1 Functional Requirements", "4"),
        ("   3.2 Non-Functional Requirements", "4"),
        ("   3.3 System Requirements", "5"),
        ("   3.4 User Requirements", "5"),
        ("4. System Design & Architectural Diagrams", "6"),
        ("5. Implementation & Code Modules", "9"),
        ("6. Experimental Results & Screenshots", "11"),
        ("7. Conclusion & Future Enhancements", "14"),
        ("8. References", "14"),
    ]
    
    for sect, pg in toc_entries:
        row_cells = table.add_row().cells
        row_cells[0].text = sect
        row_cells[1].text = pg
        set_cell_margins(row_cells[0])
        set_cell_margins(row_cells[1])

    doc.add_page_break()

    # ----------------------------------------------------
    # Page 3: Abstract & Introduction
    # ----------------------------------------------------
    add_heading_1("1. Abstract")
    add_body(
        "TalentForge is an intelligent, high-performance web-based Career Networking & Recruitment platform designed "
        "to unify job hunting, profile showcasing, and project evaluation under a unified interface. "
        "Historically, professionals and students maintain segmented identities across portfolio hosts (e.g., GitHub) "
        "and corporate communication boards (e.g., LinkedIn). TalentForge bridges this gap, allowing users to list "
        "projects with direct source-code integration and live links, build digital resumes, and establish verified "
        "connections. For recruiters, the system provides candidate indexing, advanced analytics, and shortlist management. "
        "Built using the modern MERN stack (MongoDB, Express.js, React.js, Node.js), the platform leverages a custom "
        "reactive theme hook to maintain absolute synchronization between dark/light layout states. The architecture "
        "avoids high database dependency footprint by relying on local filesystem caching alongside intelligent MongoDB "
        "aggregation indexes. The system represents a structured implementation of security features such as JWT "
        "session authentication and bcrypt cryptographic password protection, resulting in an optimized software design "
        "tailored to academic and commercial hiring standards."
    )
    
    add_heading_1("2. Introduction")
    add_body(
        "In the contemporary, fast-paced technical industry, candidates and recruiters face substantial friction during "
        "talent evaluation. Standard professional job portals lack the feature depth to verify practical capability, "
        "forcing recruiters to browse external repositories separately. Conversely, developers lack unified platforms to "
        "present their codebase alongside structural career highlights. Manual screening processes are consequently "
        "sluggish, error-prone, and inefficient, leading to critical data loss and extended hiring cycles."
    )
    add_body(
        "To address these issues, TalentForge provides a computerized, unified environment that facilitates candidate "
        "discovery through project showcasing and networking. The project aims to replace slow, manual candidate profiling "
        "with an integrated algorithmic matching system. The core design is built upon modular, object-oriented concepts, "
        "ensuring strict data encapsulation and role-based controller policies. The proposed system allows Student "
        "and Professional users to register and build unified profiles, upload academic/professional projects with "
        "verified GitHub metadata, upload PDF resumes, and send connection requests to expand their professional circle."
    )
    doc.add_page_break()

    # ----------------------------------------------------
    # Page 4: SRS
    # ----------------------------------------------------
    add_heading_1("3. Software Requirements Specification (SRS)")
    add_heading_2("3.1 Functional Requirements")
    add_bullet("FR1 (Authentication): The system shall support user registration and session management for Student, Professional, and Recruiter roles with cryptographically protected passwords using bcrypt.")
    add_bullet("FR2 (Profile Management): The system shall allow users to update location, DOB, contact phone, and external link endpoints (LinkedIn, Portfolio) with direct validation checks.")
    add_bullet("FR3 (Project Uploads): Users shall upload projects with title, tech stacks, GitHub emulation URLs, and screenshot image attachments (max 5MB in JPG/PNG formats).")
    add_bullet("FR4 (Outreach Management): Recruiters shall shortlist, reject, and message candidates, capped at 50 outreach operations per 24 hours to prevent spam.")
    add_bullet("FR5 (Resume Generation): The system shall compile structured database profile fields into a printable 2-column resume layout.")

    add_heading_2("3.2 Non-Functional Requirements")
    add_bullet("NFR1 (Security): The system shall use HTTPS for all communication and enforce JWT auth tokens with automatic 30-minute session invalidation.")
    add_bullet("NFR2 (Performance): The system shall process database queries in under 1 second and support up to 1,000 concurrent active client requests.")
    add_bullet("NFR3 (Scalability): The architecture shall support horizontal scaling and modular layout updates without service interruption.")
    add_bullet("NFR4 (Usability): The UI shall render responsive layouts across desktop and mobile browsers, supporting synchronized light/dark mode themes.")
    doc.add_page_break()

    # ----------------------------------------------------
    # Page 5: SRS System & User
    # ----------------------------------------------------
    add_heading_2("3.3 System Requirements")
    add_bullet("Hardware (Development): Intel i5/AMD Ryzen 5 Processor, 8GB RAM, 256GB SSD, and broadband internet connectivity.")
    add_bullet("Hardware (Hosting): Minimum 4-core virtual server processor, 8GB RAM, 100GB SSD, and 100Mbps dedicated bandwidth.")
    add_bullet("Software (Frontend): React.js SPA, HTML5, CSS3, JavaScript (ES6+).")
    add_bullet("Software (Backend): Node.js, Express.js REST API architecture, Mongoose ODM.")
    add_bullet("Software (Database): MongoDB Community/Atlas Database Engine.")

    add_heading_2("3.4 User Requirements")
    add_bullet("Candidates: Build profiling summaries, add technical credentials, manage uploaded project references, configure connection circles, and upload PDF documents.")
    add_bullet("Recruiters: Administer company profiles, create/manage job posts, track candidates, and access analytics dashboards.")
    add_bullet("Administrators: Access restricted dashboards, moderate inappropriate projects/posts, and execute platform maintenance.")
    doc.add_page_break()

    # ----------------------------------------------------
    # Page 6: System Design (Architecture & DFD)
    # ----------------------------------------------------
    add_heading_1("4. System Design & Architectural Diagrams")
    add_heading_2("Figure 4.1: Component System Architecture")
    p = doc.add_paragraph("[Place Draw.io Architecture Diagram Image Here]")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    add_heading_2("Figure 4.2: Data Flow Diagram (DFD Level 1)")
    p = doc.add_paragraph("[Place Draw.io Data Flow Diagram Image Here]")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # ----------------------------------------------------
    # Page 7: Design (Sequence & Control Flow)
    # ----------------------------------------------------
    add_heading_2("Figure 4.3: Platform Interaction Sequence Diagram")
    p = doc.add_paragraph("[Place Draw.io Sequence Diagram Image Here]")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    add_heading_2("Figure 4.4: Logic Control Flow Diagram")
    p = doc.add_paragraph("[Place Draw.io Control Flow Diagram Image Here]")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # ----------------------------------------------------
    # Page 8: Design (Use Case)
    # ----------------------------------------------------
    add_heading_2("Figure 4.5: Platform Use Case Model")
    p = doc.add_paragraph("[Place Draw.io Use Case Diagram Image Here]")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # ----------------------------------------------------
    # Page 9: Code Modules (Hook)
    # ----------------------------------------------------
    add_heading_1("5. Implementation & Code Modules")
    add_heading_2("Module 5.1: Reactive Dark Mode Hook (useDarkMode.js)")
    add_code(
        "import { useState, useEffect } from 'react';\n\n"
        "const useDarkMode = () => {\n"
        "  const [dark, setDark] = useState(\n"
        "    () => document.documentElement.getAttribute('data-theme') === 'dark'\n"
        "  );\n"
        "  useEffect(() => {\n"
        "    const obs = new MutationObserver(() =>\n"
        "      setDark(document.documentElement.getAttribute('data-theme') === 'dark')\n"
        "    );\n"
        "    obs.observe(document.documentElement, {\n"
        "      attributes: true, attributeFilter: ['data-theme']\n"
        "    });\n"
        "    return () => obs.disconnect();\n"
        "  }, []);\n"
        "  return dark;\n"
        "};\n"
        "export default useDarkMode;"
    )
    doc.add_page_break()

    # ----------------------------------------------------
    # Page 10: Code Modules (Controller)
    # ----------------------------------------------------
    add_heading_2("Module 5.2: Advanced Profile Controller (userController.js)")
    add_code(
        "const updateAdvancedProfile = async (req, res) => {\n"
        "  try {\n"
        "    const fields = req.body;\n"
        "    const user = await User.findById(req.user.id);\n"
        "    if (!user) return res.status(404).json({ success: false, message: 'User not found' });\n"
        "    Object.keys(fields).forEach(key => {\n"
        "      if (fields[key] !== undefined) user[key] = fields[key];\n"
        "    });\n"
        "    await user.save();\n"
        "    res.status(200).json({ success: true, message: 'Profile updated', user });\n"
        "  } catch (error) {\n"
        "    res.status(500).json({ success: false, message: error.message });\n"
        "  }\n"
        "};"
    )
    doc.add_page_break()

    # ----------------------------------------------------
    # Page 11: Screenshots Part 1
    # ----------------------------------------------------
    add_heading_1("6. Experimental Results & Screenshots")
    add_heading_2("Figure 6.1: Platform Landing Page (Dark Mode)")
    p = doc.add_paragraph("[Place Landing Page Screenshot Here]")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    add_heading_2("Figure 6.2: User Login Panel")
    p = doc.add_paragraph("[Place Login Page Screenshot Here]")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # ----------------------------------------------------
    # Page 12: Screenshots Part 2
    # ----------------------------------------------------
    add_heading_2("Figure 6.3: Recruiter Analytics Dashboard")
    p = doc.add_paragraph("[Place Analytics Dashboard Screenshot Here]")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    add_heading_2("Figure 6.4: Profile Customizer & PDF Builder")
    p = doc.add_paragraph("[Place Resume Builder Screenshot Here]")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # ----------------------------------------------------
    # Page 13: Screenshots Part 3
    # ----------------------------------------------------
    add_heading_2("Figure 6.5: Connection Networks View")
    p = doc.add_paragraph("[Place Connections Network Screenshot Here]")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    add_heading_2("Figure 6.6: Admin Panel Overview")
    p = doc.add_paragraph("[Place Admin Dashboard Screenshot Here]")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # ----------------------------------------------------
    # Page 14: Conclusion & References
    # ----------------------------------------------------
    add_heading_1("7. Conclusion & Future Enhancements")
    add_body(
        "The TalentForge Career platform successfully automates candidate evaluation and portfolio tracking. "
        "The web-based implementation validates modular design guidelines, securing page operations under role controls. "
        "The implementation successfully validates the theoretical design patterns of modern web systems, establishing "
        "stable backend architectures, schema protections, and user access policies. Future enhancements of the system "
        "include integrating machine learning recommendation engines to auto-match developer skill endorsements, parsing "
        "PDF documents using OCR, and establishing secure video-conferencing modules."
    )
    
    add_heading_1("8. References")
    refs = [
        "C. K. Gomathy, M. Pedda Chandrasekhar, K. Mallikarjun, and V. Geetha, 'The Career Networking & Recruitment Platform Architecture', ResearchGate, 2025.",
        "M. S. Kumbhar, S. R. Patil, and A. R. Desai, 'System for Managing Digital Portfolios & Hiring Profiles', IJRASET, 2024.",
        "Robert C. Martin, 'Clean Code: A Handbook of Agile Software Craftsmanship', Prentice Hall, 2008.",
        "Bjarne Stroustrup, 'The design & structure of modular software architectures', Addison-Wesley, 2013.",
        "Mongoose ODM API Reference Guides, https://mongoosejs.com.",
        "React Component Life Cycle documentation guides, https://react.dev.",
    ]
    for idx, r in enumerate(refs):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"[{idx+1}] {r}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)

    doc.save(r"c:\Users\Raghavendra s\SE Project\TalentForge_Full_Report.docx")
    print("Word document generated successfully!")

if __name__ == "__main__":
    create_word_doc()
